# -*- coding: utf-8 -*-
"""
Created on Tue May  5 19:15:19 2020

@author: visha
"""
import HtmlTestRunner
import unittest
from selenium import webdriver
from selenium.webdriver.common.keys import Keys

l = [["search_google", "google", "a",], ["search_facebook", "facebook", "b"], ["search_twitter", "twitter", "b"]]

class TestSequense(unittest.TestCase):
    @classmethod
    def setUp(self):
        # create a new Firefox session
        self.driver = webdriver.Chrome(r'chromedriver.exe')
        self.driver.implicitly_wait(30)
        self.driver.maximize_window()
        print(self.driver.title)
        # navigate to the application home page
        self.driver.get("http://www.google.com/")

    @classmethod
    def tearDown(self):
        # close the browser window
        self.driver.quit()

imageList = []
def test_generator(a,b):
    def test_search_by_text(self):
        # get the search textbox
        self.search_field = self.driver.find_element_by_name("q")

        # enter search keyword and submit
        self.search_field.send_keys(b)
        self.search_field.submit()

        #get the list of elements which are displayed after the search
        #currently on result page usingfind_elements_by_class_namemethod

        lists = self.driver.find_elements_by_class_name("r")
        print("taking screenshot " + a)
        self.driver.save_screenshot(a + ".png")
        self.driver.implicitly_wait(5)
        imageList.append(a + ".png")
        self.assertEqual(11, len(lists))
        
    return test_search_by_text

from docx import Document
from docx.shared import Inches

def generateReport():
    #file_ref = open("addImage.docx","rb")
    doc = Document("addImage.docx")
    i=0
    for image in imageList:
        tables = doc.tables
        p = tables[i].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(image,width=Inches(4.0), height=Inches(.7))
        i=i+1
    doc.save("addImage.docx")

if __name__ == '__main__':
    for t in l:
        test_name = 'test_%s' % t[0]
        test = test_generator(test_name, t[1])
        setattr(TestSequense, test_name, test)
    unittest.main(testRunner=HtmlTestRunner.HTMLTestRunner())
    generateReport()